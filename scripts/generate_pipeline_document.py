from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "docs"
OUTPUT_PATH = OUTPUT_DIR / "Supplier_Redaction_Pipeline_Technical_Documentation.docx"


def add_title(document: Document, title: str, subtitle: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(title)
    run.bold = True
    run.font.size = Pt(22)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(subtitle)
    run.italic = True
    run.font.size = Pt(11)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    document.add_heading(text, level=level)


def add_para(document: Document, text: str, style: str | None = None) -> None:
    paragraph = document.add_paragraph(style=style)
    paragraph.add_run(text)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        add_para(document, item, "List Bullet")


def add_numbered(document: Document, items: list[str]) -> None:
    for item in items:
        add_para(document, item, "List Number")


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value


def add_code_block(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def build_document() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    add_title(
        document,
        "Supplier Redaction Workspace",
        "Technical Architecture, Pipeline Design, and Operational Documentation",
    )

    add_para(
        document,
        "This document describes the current supplier redaction platform as implemented in this repository. "
        "It reflects the latest architecture where Streamlit is the business-facing presentation layer, FastAPI is the backend service boundary, and the redaction engine is a shared reusable core. "
        "The document is written for engineers, technical reviewers, deployment owners, and future maintainers who need to understand how documents move through the system, why each dependency exists, and where the production scaling boundaries are.",
    )

    add_heading(document, "1. Executive Summary", 1)
    add_para(
        document,
        "The system redacts supplier-identifying information from tender and contract documents. It supports PDF and DOCX input, produces redacted PDF or DOCX output, and records redaction metadata for operator review. "
        "The application is currently optimized for local Windows execution but is structured for future deployment by separating UI concerns from backend processing.",
    )
    add_bullets(
        document,
        [
            "Streamlit provides the user workflow: upload, preview, options, run, summary, and download.",
            "FastAPI provides the service boundary: health check, redaction request, and secure output download.",
            "The redaction engine performs document parsing, OCR coordination, supplier-name resolution, text redaction, image redaction, and output creation.",
            "Manual supplier names remain the most controlled source of truth; optional heuristics and LLM extraction are discovery aids.",
            "Runtime files are isolated under data/ so source code, uploads, outputs, and temporary files are separated.",
        ],
    )

    add_heading(document, "2. Current Runtime Architecture", 1)
    add_para(
        document,
        "The production-oriented runtime shape is a two-layer application: Streamlit calls FastAPI over HTTP, and FastAPI delegates to the shared redaction engine. "
        "This keeps the UI thin and makes the backend callable by future integrations without depending on Streamlit internals.",
    )
    add_code_block(
        document,
        "Browser\n"
        "  |\n"
        "  v\n"
        "Streamlit UI (streamlit_app.py, port 8501)\n"
        "  |  HTTP multipart request via httpx\n"
        "  v\n"
        "FastAPI Backend (app/main.py, port 8000)\n"
        "  |  RedactionOptions + temporary upload path\n"
        "  v\n"
        "Shared Redaction Engine (app/redactors.py)\n"
        "  |-- Supplier parsing/detection (app/detection.py)\n"
        "  |-- Tesseract OCR (app/ocr.py)\n"
        "  |-- Optional Azure OCR (app/document_intelligence.py)\n"
        "  |-- PDF/DOCX redaction and output writing\n"
        "  v\n"
        "data/outputs/<redacted-document>\n"
        "  ^\n"
        "  |  GET /download?path=<output_path>\n"
        "Streamlit download button",
    )

    add_heading(document, "3. Architectural Principles", 1)
    add_table(
        document,
        ["Principle", "Implementation", "Reason"],
        [
            [
                "Clear service boundary",
                "Streamlit calls FastAPI instead of importing the redaction engine directly.",
                "The same backend path can be used locally, in production, and by future integrations.",
            ],
            [
                "Reusable core engine",
                "PDF/DOCX processing lives in app/redactors.py and is independent of Streamlit.",
                "The pipeline can be invoked by API, CLI, or tests without duplicating logic.",
            ],
            [
                "Conservative redaction behavior",
                "Manual names, exact/native matches, OCR token matches, and explicit image options are used.",
                "Redaction is high impact; deterministic behavior is preferred over aggressive guessing.",
            ],
            [
                "Config outside code",
                "pydantic-settings reads defaults and optional .env overrides.",
                "Deployment-specific values such as API URL, OCR, OpenAI, and Azure settings do not require code edits.",
            ],
            [
                "Runtime isolation",
                "Uploads, outputs, Streamlit home, and temp files live under data/.",
                "Keeps repository source clean and simplifies deployment/storage planning.",
            ],
        ],
    )

    add_heading(document, "4. End-to-End Processing Flow", 1)
    add_numbered(
        document,
        [
            "The user starts FastAPI with run_api.ps1 and Streamlit with run_ui.ps1.",
            "The user opens the Streamlit UI and uploads a PDF or DOCX file.",
            "Streamlit previews the first PDF page in memory using PyMuPDF; DOCX files are accepted without visual preview.",
            "The user enters supplier names and selects discovery, image redaction, OCR, and replacement-label options.",
            "Streamlit posts the file and options to POST /redact using httpx.",
            "FastAPI validates the extension, writes the upload to a temporary per-request folder, and creates RedactionOptions.",
            "The redaction dispatcher routes the file to the PDF or DOCX processing path.",
            "The pipeline extracts native text and optionally uses OCR/LLM/heuristic sources to resolve supplier names.",
            "The pipeline applies text and image redactions, writes the output under data/outputs, and returns RedactionResult metadata.",
            "FastAPI adds a download_url to the response and schedules cleanup of the temporary upload folder.",
            "Streamlit retrieves the generated file from the download_url and presents the download button, metrics, warnings, names used, sources, and redaction details.",
        ],
    )

    add_heading(document, "5. User-Facing Controls and Backend Mapping", 1)
    add_table(
        document,
        ["UI Control", "Backend Field", "Behavior"],
        [
            [
                "Known supplier names",
                "supplier_names",
                "Parsed by shared logic. Supports new lines, semicolons, and separator commas while preserving legal suffix commas such as Company, LLC.",
            ],
            [
                "Manual names only",
                "use_llm_extraction=false, detect_supplier_names=false",
                "Uses only operator-provided names and their generated legal-name variants.",
            ],
            [
                "Use LLM extraction",
                "use_llm_extraction=true",
                "Requests OpenAI-based supplier/legal-party extraction when OPENAI_API_KEY is configured.",
            ],
            [
                "Use LLM + heuristic detection",
                "use_llm_extraction=true, detect_supplier_names=true",
                "Combines manual names, LLM names, and local legal-entity pattern candidates.",
            ],
            [
                "Header/footer logos",
                "redact_header_footer_images=true",
                "Redacts embedded PDF images near the top or bottom of pages, usually logos or letterhead.",
            ],
            [
                "All images",
                "redact_all_images=true",
                "Redacts every embedded image object in supported documents.",
            ],
            [
                "Do not redact images",
                "redact_all_images=false, redact_header_footer_images=false",
                "Leaves image objects unchanged; text redaction still runs.",
            ],
            [
                "Auto OCR",
                "use_ocr=null",
                "Uses configured OCR policy. Default auto mode runs OCR only when native PDF text is sparse.",
            ],
            [
                "Force OCR",
                "use_ocr=true",
                "Runs Tesseract even when selectable PDF text is present.",
            ],
            [
                "Disable OCR",
                "use_ocr=false",
                "Skips Tesseract OCR completely.",
            ],
            [
                "Replacement label",
                "replacement_text",
                "Text written into redaction annotations or DOCX replacement text. Defaults to REDACTED.",
            ],
        ],
    )

    add_heading(document, "6. Supplier Name Handling", 1)
    add_para(
        document,
        "Supplier names are normalized, parsed, expanded, merged, and deduplicated before redaction. This prevents the system from depending on only a single exact phrase when legal documents split parties across lines or vary legal suffix formatting.",
    )
    add_bullets(
        document,
        [
            "Manual entries are parsed from new lines, semicolons, and separator commas.",
            "Legal commas in suffixes such as LLC are preserved as part of the supplier name.",
            "Combined party names can be split on connectors such as and, &, and |.",
            "Legal suffix variants are generated for common forms such as LLC and L.L.C.",
            "Generated names are deduplicated while preserving useful order.",
            "The pipeline remains conservative: it will not treat a different company as the same supplier unless the name is supplied or discovered by enabled extraction.",
        ],
    )

    add_heading(document, "7. PDF Redaction Path", 1)
    add_numbered(
        document,
        [
            "Open the PDF with PyMuPDF.",
            "Extract native selectable text from every page.",
            "Run Azure Document Intelligence only when endpoint/key settings are configured.",
            "Decide whether Tesseract OCR should run based on request override and OCR settings.",
            "Combine native text, Azure text, and OCR text for supplier-name resolution.",
            "Resolve supplier names from manual input, optional heuristics, and optional OpenAI extraction.",
            "Search native PDF text with PyMuPDF and add redaction annotations.",
            "Find OCR token-sequence matches and redact the corresponding OCR bounding boxes.",
            "Select image rectangles based on header/footer or all-image settings.",
            "Apply redactions physically and save a new optimized PDF under data/outputs.",
        ],
    )

    add_heading(document, "8. DOCX Redaction Path", 1)
    add_numbered(
        document,
        [
            "Open the document with python-docx.",
            "Collect paragraph text for supplier-name resolution.",
            "Resolve names from manual input, optional heuristics, and optional OpenAI extraction.",
            "Walk top-level paragraphs and table cell paragraphs.",
            "Replace matching text inside Word runs.",
            "Save the redacted DOCX under data/outputs.",
            "When image redaction is enabled, remove embedded media files and relationships from the DOCX zip package.",
        ],
    )

    add_heading(document, "9. Component Responsibilities", 1)
    add_table(
        document,
        ["File", "Responsibility", "Architectural Notes"],
        [
            [
                "streamlit_app.py",
                "Business UI, preview, option collection, API calls, result display, and download button.",
                "Presentation layer only; it does not perform redaction directly.",
            ],
            [
                "app/main.py",
                "FastAPI endpoints: /health, /redact, and /download.",
                "Owns HTTP validation, temporary upload handling, download safety checks, and response contract.",
            ],
            [
                "app/models.py",
                "Pydantic models for options, hits, and results.",
                "Shared contract across API, CLI, Streamlit, and redaction engine.",
            ],
            [
                "app/config.py",
                "Runtime settings loaded from defaults and .env.",
                "Centralizes API URL, storage paths, OCR, OpenAI, Azure, and image defaults.",
            ],
            [
                "app/redactors.py",
                "Core PDF/DOCX redaction engine.",
                "Contains the file-type dispatcher and all document mutation logic.",
            ],
            [
                "app/detection.py",
                "Supplier parsing, normalization, alias expansion, heuristic extraction, and OpenAI extraction.",
                "Keeps name-resolution behavior shared across UI/API/CLI.",
            ],
            [
                "app/ocr.py",
                "Local Tesseract OCR and OCR coordinate mapping.",
                "Converts page images to OCR words and maps OCR hits back to PDF rectangles.",
            ],
            [
                "app/document_intelligence.py",
                "Optional Azure Document Intelligence text extraction.",
                "Loaded only when endpoint/key settings are present.",
            ],
            [
                "app/cli.py",
                "Local command-line execution path.",
                "Useful for developer checks and batch-style runs without Streamlit.",
            ],
            [
                "scripts/smoke_test_pipeline.py",
                "End-to-end verification script.",
                "Checks health, upload, redaction, download, warnings, and API/direct consistency.",
            ],
            [
                "scripts/generate_pipeline_document.py",
                "Generates this technical documentation.",
                "Keeps architecture documentation reproducible from source.",
            ],
        ],
    )

    add_heading(document, "10. Dependency Rationale", 1)
    add_para(
        document,
        "The dependency set is intentionally small and aligned to one of five concerns: user interface, backend service boundary, document manipulation, OCR/text extraction, and optional AI-assisted discovery.",
    )
    add_table(
        document,
        ["Package", "Purpose in this project", "Why it is used"],
        [
            ["fastapi", "Backend API for upload, redaction, health checks, and downloads.", "Provides a production-ready service boundary behind the Streamlit UI."],
            ["uvicorn[standard]", "ASGI server for running the FastAPI app.", "Needed to serve the backend locally and in deployment."],
            ["httpx", "HTTP client used by Streamlit to call FastAPI.", "Lets the UI use the same backend path that production integrations will use."],
            ["python-multipart", "Multipart form/file parsing for FastAPI.", "Required for POST /redact to receive uploaded PDF/DOCX files and options."],
            ["pydantic-settings", "Runtime configuration from defaults and .env.", "Keeps paths, OCR, OpenAI, Azure, and API URL settings centralized."],
            ["streamlit", "Business-facing web interface.", "Provides quick upload, preview, option selection, result summary, and download UI."],
            ["PyMuPDF", "PDF reading, preview rendering, text extraction, redaction, image detection, and PDF saving.", "Provides the core PDF operations needed by the redaction pipeline."],
            ["python-docx", "DOCX reading, text replacement, media removal, and saving.", "Enables Word document redaction without converting files to another format."],
            ["Pillow", "Image object handling for previews and OCR input.", "Bridges rendered PDF pages and image data used by Streamlit/Tesseract."],
            ["pytesseract", "Python wrapper for Tesseract OCR.", "Finds text in scanned or image-based PDFs when OCR is enabled."],
            ["openai", "Optional LLM-based supplier/vendor/legal-party extraction.", "Helps identify supplier names when manual input is incomplete."],
            ["azure-ai-documentintelligence", "Optional Azure Document Intelligence OCR/text extraction.", "Supports cloud OCR when Azure endpoint/key settings are configured."],
        ],
    )

    add_heading(document, "11. Configuration Reference", 1)
    add_table(
        document,
        ["Setting", "Default", "Meaning"],
        [
            ["APP_NAME", "Supplier Redaction API", "FastAPI title."],
            ["DATA_DIR", "data", "Base application data directory."],
            ["OUTPUT_DIR", "data/outputs", "Where redacted files are written."],
            ["UPLOAD_DIR", "data/uploads", "Where uploaded files are temporarily staged."],
            ["REDACTION_API_URL", "http://127.0.0.1:8000", "Backend API base URL used by Streamlit."],
            ["REDACT_ALL_IMAGES", "false", "Default policy for redacting all PDF images."],
            ["REDACT_HEADER_FOOTER_IMAGES", "true", "Default policy for redacting likely header/footer logos."],
            ["USE_OPENAI_EXTRACTION", "false", "Default LLM extraction behavior when request-level value is null."],
            ["OPENAI_API_KEY", "empty", "Enables OpenAI extraction when present and requested."],
            ["OPENAI_MODEL", "gpt-5.4-mini", "OpenAI model used by extract_with_openai()."],
            ["USE_TESSERACT_OCR", "true", "Global enable/disable for local OCR."],
            ["TESSERACT_CMD", "None", "Optional full path to tesseract.exe."],
            ["OCR_MODE", "auto", "off disables OCR, always forces OCR, auto runs OCR when native text is sparse."],
            ["OCR_DPI", "220", "Resolution used when rendering PDF pages for OCR."],
            ["OCR_MIN_CONFIDENCE", "45", "Minimum accepted Tesseract confidence for OCR words."],
            ["NATIVE_TEXT_MIN_CHARS", "80", "Native text threshold used by OCR auto mode."],
            ["AZURE_DOCUMENT_INTELLIGENCE_ENDPOINT", "None", "Azure endpoint for optional cloud OCR."],
            ["AZURE_DOCUMENT_INTELLIGENCE_KEY", "None", "Azure key for optional cloud OCR."],
        ],
    )

    add_heading(document, "12. API Contract", 1)
    add_para(document, "FastAPI is the backend contract for Streamlit and future integrations.")
    add_table(
        document,
        ["Endpoint", "Purpose", "Important Behavior"],
        [
            ["GET /health", "Liveness check.", "Returns {'status': 'ok'} when the backend is running."],
            ["POST /redact", "Main redaction request.", "Accepts multipart file/options, validates file type, runs the redaction engine, and returns RedactionResult."],
            ["GET /download", "Output retrieval.", "Allows downloads only from inside the configured output directory."],
        ],
    )
    add_table(
        document,
        ["POST /redact Field", "Type", "Purpose"],
        [
            ["file", "UploadFile", "PDF or DOCX file to redact."],
            ["supplier_names", "string", "Supplier names separated by new lines, semicolons, or separator commas."],
            ["detect_supplier_names", "boolean", "Enables regex-based organization detection."],
            ["use_llm_extraction", "boolean/null", "Forces or disables OpenAI extraction for this request. Null uses settings."],
            ["use_ocr", "boolean/null", "Forces or disables Tesseract OCR for this request. Null uses settings."],
            ["redact_all_images", "boolean", "Redacts every image in the file when supported."],
            ["redact_header_footer_images", "boolean", "Redacts likely logos in header/footer areas."],
            ["replacement_text", "string", "Replacement label written over redacted text."],
        ],
    )

    add_heading(document, "13. Result Metadata", 1)
    add_table(
        document,
        ["Field", "Meaning"],
        [
            ["input_path", "Path of the temporary uploaded input or CLI input file."],
            ["output_path", "Path of the generated redacted file under data/outputs."],
            ["download_url", "API-relative URL used by Streamlit to download the generated file."],
            ["media_type", "MIME type of the generated output."],
            ["hits", "List of text/image redaction events."],
            ["supplier_names", "Final merged supplier-name list used for redaction."],
            ["extraction_sources", "Labels showing which extraction methods contributed."],
            ["warnings", "Operational warnings such as missing Tesseract or missing OPENAI_API_KEY."],
        ],
    )

    add_heading(document, "14. Runtime Data Management", 1)
    add_table(
        document,
        ["Location", "Purpose", "Retention Notes"],
        [
            ["data/uploads", "Temporary uploaded input copies.", "API requests use per-request folders and schedule cleanup after processing."],
            ["data/outputs", "Generated redacted files.", "Files remain available for download and audit until retention cleanup is added."],
            ["data/streamlit_tmp", "Streamlit temporary files.", "Kept inside the project by run_ui.ps1."],
            ["data/streamlit_home", "Streamlit home/config location.", "Avoids writes to C:\\Users\\<user>\\.streamlit."],
            ["contracts", "Optional local sample documents.", "Ignored by Git so private contracts are not published."],
            ["docs", "Generated technical documentation.", "The DOCX is generated by scripts/generate_pipeline_document.py."],
        ],
    )

    add_heading(document, "15. Validation and Smoke Testing", 1)
    add_para(document, "The smoke test validates the deployed-style API path and compares it against direct engine execution for the sample contract.")
    add_bullets(
        document,
        [
            "Verifies GET /health.",
            "Uploads a sample PDF through POST /redact.",
            "Checks that the output file is created.",
            "Downloads the result through GET /download.",
            "Compares API supplier names and hit count with direct redaction engine execution.",
            "Runs an OCR/LLM-requested path and expects warnings rather than crashes when external prerequisites are unavailable.",
        ],
    )
    add_code_block(
        document,
        "$env:PYTHONDONTWRITEBYTECODE='1'\n"
        ".\\.venv\\Scripts\\python.exe -B scripts\\smoke_test_pipeline.py",
    )

    add_heading(document, "16. Security and Compliance Considerations", 1)
    add_bullets(
        document,
        [
            "Downloaded files are restricted to the configured output directory to prevent arbitrary file reads.",
            "Uploaded filenames are sanitized to the basename and staged under per-request upload folders.",
            "Secrets such as OPENAI_API_KEY and Azure keys belong in .env or deployment secret stores, not source control.",
            "Redacted outputs may still be sensitive; production deployments should define retention, access control, and audit logging policies.",
            "OCR and LLM extraction can process sensitive text. Production usage should confirm data handling requirements for each external service.",
            "The current local deployment has no authentication. Public deployment must add authentication and authorization before exposing upload/download endpoints.",
        ],
    )

    add_heading(document, "17. Known Limitations", 1)
    add_bullets(
        document,
        [
            "Tesseract must be installed separately; pytesseract is only the Python wrapper.",
            "OpenAI extraction only runs when OPENAI_API_KEY is configured and extraction is requested.",
            "Heuristic detection is optional because contracts often mention organizations that are not suppliers.",
            "DOCX text split across multiple Word runs may not always be redacted as a continuous phrase.",
            "OCR matching is conservative and token-sequence based; very noisy scans may require future fuzzy matching.",
            "Image redaction applies to embedded image objects. Normal selectable header/footer text is handled by text redaction, not image settings.",
            "Long-running OCR-heavy workloads are currently synchronous HTTP requests; production scale may require background jobs.",
        ],
    )

    add_heading(document, "18. Production Readiness Roadmap", 1)
    add_table(
        document,
        ["Area", "Recommended Next Step", "Reason"],
        [
            ["Authentication", "Add login/session or identity-provider integration.", "Prevent unauthorized upload and download access."],
            ["Storage", "Move outputs to managed object storage with signed download URLs.", "Avoid local disk coupling when API/UI are deployed separately."],
            ["Background jobs", "Introduce a queue for OCR-heavy or large documents.", "Avoid HTTP timeouts and improve scalability."],
            ["Audit logging", "Persist request metadata, redaction hits, warnings, and operator decisions.", "Supports traceability and compliance review."],
            ["Candidate review", "Add a confirmation screen for LLM/heuristic names before redaction.", "Reduces over-redaction risk from automated discovery."],
            ["Retention", "Add scheduled cleanup for uploads and old outputs.", "Controls storage growth and sensitive-data exposure."],
            ["Observability", "Add structured logs, metrics, and error tracking.", "Improves production support and incident response."],
            ["Testing", "Add unit tests for parsing, alias expansion, OCR matching, API options, and DOCX edge cases.", "Protects behavior as the pipeline evolves."],
        ],
    )

    add_heading(document, "19. Operating Commands", 1)
    add_table(
        document,
        ["Task", "Command"],
        [
            ["Create environment", "python -m venv .venv"],
            ["Install dependencies", ".\\.venv\\Scripts\\Activate.ps1; pip install -r requirements.txt"],
            ["Start API", ".\\run_api.ps1"],
            ["Start Streamlit UI", ".\\run_ui.ps1"],
            ["Run smoke test", "$env:PYTHONDONTWRITEBYTECODE='1'; .\\.venv\\Scripts\\python.exe -B scripts\\smoke_test_pipeline.py"],
            ["Regenerate this document", ".\\.venv\\Scripts\\python.exe -B scripts\\generate_pipeline_document.py"],
        ],
    )

    add_heading(document, "20. Summary", 1)
    add_para(
        document,
        "The pipeline is now structured in a production-aligned way: Streamlit owns the business workflow, FastAPI owns the backend service boundary, and the redaction engine owns deterministic document processing. "
        "This architecture keeps the current local workflow simple while preserving a clear path to hosted deployment, API integrations, background processing, and stronger operational controls.",
    )

    document.save(OUTPUT_PATH)


if __name__ == "__main__":
    build_document()
    print(OUTPUT_PATH)
