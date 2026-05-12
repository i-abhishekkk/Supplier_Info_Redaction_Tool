import os
import re
import zipfile
from pathlib import Path
from uuid import uuid4

import fitz

from app.config import get_settings
from app.detection import (
    extract_supplier_candidates,
    extract_with_openai,
    merge_supplier_names,
)
from app.document_intelligence import extract_text_with_document_intelligence
from app.models import RedactionHit, RedactionOptions, RedactionResult
from app.ocr import find_ocr_name_rects, find_ocr_sensitive_rects, run_pdf_ocr
from app.sensitive import SensitiveField, find_sensitive_text_matches, replace_sensitive_text


PDF_MEDIA_TYPE = "application/pdf"
DOCX_MEDIA_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


async def redact_document(input_path: Path, options: RedactionOptions) -> RedactionResult:
    suffix = input_path.suffix.lower()
    if suffix == ".pdf":
        return await redact_pdf(input_path, options)
    if suffix == ".docx":
        return await redact_docx(input_path, options)
    raise ValueError(f"Unsupported file type: {input_path.suffix}")


async def redact_pdf(input_path: Path, options: RedactionOptions) -> RedactionResult:
    settings = get_settings()
    output_path = settings.output_dir / f"{input_path.stem}.redacted.{uuid4().hex[:8]}.pdf"
    extraction_sources: list[str] = []
    warnings: list[str] = []

    doc = fitz.open(input_path)
    native_text = "\n".join(page.get_text("text") for page in doc)
    doc.close()
    if native_text.strip():
        extraction_sources.append("pdf-native-text")

    di_text = await extract_text_with_document_intelligence(
        input_path,
        settings.azure_document_intelligence_endpoint,
        settings.azure_document_intelligence_key,
    )
    if di_text:
        extraction_sources.append("azure-document-intelligence")

    ocr_result = None
    if _should_run_ocr(native_text, options):
        ocr_result = run_pdf_ocr(input_path)
        warnings.extend(ocr_result.warnings)
        if ocr_result.text.strip():
            extraction_sources.append("tesseract-ocr")

    text_parts = [native_text]
    if di_text:
        text_parts.append(di_text)
    if ocr_result and ocr_result.text:
        text_parts.append(ocr_result.text)
    analysis_text = "\n".join(text_parts)
    supplier_names, name_sources, name_warnings = await _resolve_supplier_names(analysis_text, options)
    extraction_sources.extend(name_sources)
    if options.sensitive_fields:
        extraction_sources.append("sensitive-pattern-detection")
    warnings.extend(name_warnings)

    hits: list[RedactionHit] = []
    ocr_rects = find_ocr_name_rects(ocr_result.pages, supplier_names) if ocr_result else []
    ocr_sensitive_rects = find_ocr_sensitive_rects(ocr_result.pages, options.sensitive_fields) if ocr_result else []
    doc = fitz.open(input_path)
    ocr_rects_by_page: dict[int, list[tuple[fitz.Rect, str]]] = {}
    for page_number, rect, name in ocr_rects:
        ocr_rects_by_page.setdefault(page_number, []).append((rect, name))
    ocr_sensitive_rects_by_page: dict[int, list[tuple[fitz.Rect, str]]] = {}
    for page_number, rect, label in ocr_sensitive_rects:
        ocr_sensitive_rects_by_page.setdefault(page_number, []).append((rect, label))

    for page_index, page in enumerate(doc):
        redacted_rects: list[fitz.Rect] = []
        page_hits = _redact_page_text(page, supplier_names, options.replacement_text, redacted_rects)
        hits.extend(RedactionHit(kind="text", page=page_index + 1, value=value, reason="supplier-name-match") for value in page_hits)
        sensitive_hits = _redact_page_sensitive_text(page, options.sensitive_fields, options.replacement_text, redacted_rects)
        hits.extend(
            RedactionHit(kind="text", page=page_index + 1, value=value, reason="sensitive-info-match")
            for value in sensitive_hits
        )
        for rect, name in ocr_rects_by_page.get(page_index + 1, []):
            if any(rect.intersects(existing) for existing in redacted_rects):
                continue
            page.add_redact_annot(rect, text=options.replacement_text, fill=(0, 0, 0), text_color=(1, 1, 1))
            redacted_rects.append(rect)
            hits.append(
                RedactionHit(kind="text", page=page_index + 1, value=name, reason="ocr-supplier-name-match")
            )
        for rect, label in ocr_sensitive_rects_by_page.get(page_index + 1, []):
            if any(rect.intersects(existing) for existing in redacted_rects):
                continue
            page.add_redact_annot(rect, text=options.replacement_text, fill=(0, 0, 0), text_color=(1, 1, 1))
            redacted_rects.append(rect)
            hits.append(RedactionHit(kind="text", page=page_index + 1, value=label, reason="ocr-sensitive-info-match"))

        for rect, reason in _image_redaction_rects(page, options):
            page.add_redact_annot(rect, fill=(0, 0, 0))
            hits.append(RedactionHit(kind="image", page=page_index + 1, reason=reason))

        page.apply_redactions(images=fitz.PDF_REDACT_IMAGE_PIXELS)

    doc.save(output_path, garbage=4, deflate=True)
    doc.close()
    return RedactionResult(
        input_path=input_path,
        output_path=output_path,
        media_type=PDF_MEDIA_TYPE,
        hits=hits,
        supplier_names=supplier_names,
        sensitive_fields=options.sensitive_fields,
        extraction_sources=_dedupe(extraction_sources),
        warnings=warnings,
    )


async def redact_docx(input_path: Path, options: RedactionOptions) -> RedactionResult:
    from docx import Document

    settings = get_settings()
    output_path = settings.output_dir / f"{input_path.stem}.redacted.{uuid4().hex[:8]}.docx"
    doc = Document(input_path)
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    supplier_names, name_sources, name_warnings = await _resolve_supplier_names(text, options)
    extraction_sources = ["docx-text", *name_sources]
    if options.sensitive_fields:
        extraction_sources.append("sensitive-pattern-detection")
    hits: list[RedactionHit] = []

    for paragraph in doc.paragraphs:
        hits.extend(_redact_docx_paragraph(paragraph, supplier_names, options.replacement_text))
        hits.extend(_redact_docx_paragraph_sensitive(paragraph, options.sensitive_fields, options.replacement_text))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    hits.extend(_redact_docx_paragraph(paragraph, supplier_names, options.replacement_text))
                    hits.extend(_redact_docx_paragraph_sensitive(paragraph, options.sensitive_fields, options.replacement_text))

    doc.save(output_path)
    if options.redact_all_images or options.redact_header_footer_images:
        _remove_docx_media(output_path)
        hits.append(RedactionHit(kind="image", reason="docx-media-removed"))

    return RedactionResult(
        input_path=input_path,
        output_path=output_path,
        media_type=DOCX_MEDIA_TYPE,
        hits=hits,
        supplier_names=supplier_names,
        sensitive_fields=options.sensitive_fields,
        extraction_sources=extraction_sources,
        warnings=name_warnings,
    )


async def _resolve_supplier_names(text: str, options: RedactionOptions) -> tuple[list[str], list[str], list[str]]:
    settings = get_settings()
    sources: list[str] = []
    warnings: list[str] = []
    heuristic_names = extract_supplier_candidates(text) if options.detect_supplier_names else []
    if heuristic_names:
        sources.append("heuristic-entity-detection")
    llm_names = []
    use_llm = settings.use_openai_extraction if options.use_llm_extraction is None else options.use_llm_extraction
    if use_llm:
        if os.getenv("OPENAI_API_KEY"):
            llm_names = await extract_with_openai(text, settings.openai_model)
            if llm_names:
                sources.append("openai-llm-extraction")
        else:
            warnings.append("LLM extraction was requested, but OPENAI_API_KEY is not set.")
    if options.supplier_names:
        sources.append("user-supplied-names")
    return merge_supplier_names(options.supplier_names, heuristic_names, llm_names), sources, warnings


def _should_run_ocr(native_text: str, options: RedactionOptions) -> bool:
    settings = get_settings()
    if options.use_ocr is not None:
        return options.use_ocr
    if not settings.use_tesseract_ocr or settings.ocr_mode.lower() == "off":
        return False
    if settings.ocr_mode.lower() == "always":
        return True
    return len(native_text.strip()) < settings.native_text_min_chars


def _redact_page_text(
    page: fitz.Page,
    supplier_names: list[str],
    replacement_text: str,
    redacted_rects: list[fitz.Rect] | None = None,
) -> list[str]:
    hits: list[str] = []
    redacted_rects = redacted_rects if redacted_rects is not None else []
    names = sorted(supplier_names, key=len, reverse=True)
    for name in names:
        flags = fitz.TEXT_DEHYPHENATE | fitz.TEXT_PRESERVE_WHITESPACE
        for rect in page.search_for(name, flags=flags):
            if any(rect.intersects(existing) for existing in redacted_rects):
                continue
            page.add_redact_annot(rect, text=replacement_text, fill=(0, 0, 0), text_color=(1, 1, 1))
            redacted_rects.append(rect)
            hits.append(name)
    return hits


def _redact_page_sensitive_text(
    page: fitz.Page,
    fields: list[SensitiveField],
    replacement_text: str,
    redacted_rects: list[fitz.Rect] | None = None,
) -> list[str]:
    if not fields:
        return []

    hits: list[str] = []
    redacted_rects = redacted_rects if redacted_rects is not None else []
    flags = fitz.TEXT_DEHYPHENATE | fitz.TEXT_PRESERVE_WHITESPACE
    page_text = page.get_text("text")
    for match in sorted(find_sensitive_text_matches(page_text, fields), key=lambda item: len(item.value), reverse=True):
        rects = page.search_for(match.value, flags=flags)
        if not rects and match.field == "address":
            rects = _address_block_rects(page, match.value)
        for rect in rects:
            if any(rect.intersects(existing) for existing in redacted_rects):
                continue
            page.add_redact_annot(rect, text=replacement_text, fill=(0, 0, 0), text_color=(1, 1, 1))
            redacted_rects.append(rect)
            hits.append(match.label)
    return hits


def _address_block_rects(page: fitz.Page, value: str) -> list[fitz.Rect]:
    target = re.sub(r"\s+", " ", value).lower()
    rects: list[fitz.Rect] = []
    for block in page.get_text("blocks"):
        if len(block) < 5:
            continue
        text = re.sub(r"\s+", " ", str(block[4])).lower()
        if target and target in text:
            rects.append(fitz.Rect(block[:4]))
    return rects


def _image_redaction_rects(page: fitz.Page, options: RedactionOptions) -> list[tuple[fitz.Rect, str]]:
    settings = get_settings()
    redact_all = options.redact_all_images or settings.redact_all_images
    redact_header_footer = options.redact_header_footer_images or settings.redact_header_footer_images
    if not redact_all and not redact_header_footer:
        return []

    rects: list[tuple[fitz.Rect, str]] = []
    page_height = page.rect.height
    header_cutoff = page_height * 0.18
    footer_cutoff = page_height * 0.82
    seen: set[tuple[float, float, float, float]] = set()

    for image_info in page.get_images(full=True):
        xref = image_info[0]
        for rect in page.get_image_rects(xref):
            key = (round(rect.x0, 2), round(rect.y0, 2), round(rect.x1, 2), round(rect.y1, 2))
            if key in seen:
                continue
            seen.add(key)
            if redact_all:
                rects.append((rect, "all-images"))
            elif rect.y0 <= header_cutoff or rect.y1 >= footer_cutoff:
                rects.append((rect, "header-footer-image"))
    return rects


def _redact_docx_paragraph(paragraph, supplier_names: list[str], replacement_text: str) -> list[RedactionHit]:
    hits: list[RedactionHit] = []
    for run in paragraph.runs:
        original = run.text
        redacted = original
        for name in supplier_names:
            redacted, count = re.subn(re.escape(name), replacement_text, redacted, flags=re.IGNORECASE)
            for _ in range(count):
                hits.append(RedactionHit(kind="text", value=name, reason="supplier-name-match"))
        if redacted != original:
            run.text = redacted
    return hits


def _redact_docx_paragraph_sensitive(
    paragraph,
    fields: list[SensitiveField],
    replacement_text: str,
) -> list[RedactionHit]:
    hits: list[RedactionHit] = []
    if not fields:
        return hits
    for run in paragraph.runs:
        original = run.text
        redacted, labels = replace_sensitive_text(original, fields, replacement_text)
        if labels:
            run.text = redacted
            hits.extend(RedactionHit(kind="text", value=label, reason="sensitive-info-match") for label in labels)
    return hits


def _remove_docx_media(path: Path) -> None:
    temp_path = path.with_suffix(".tmp.docx")
    media_prefix = "word/media/"
    relationship_media_pattern = re.compile(r'<Relationship[^>]+Target="media/[^"]+"[^>]*/>')
    with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(temp_path, "w", zipfile.ZIP_DEFLATED) as target:
        for item in source.infolist():
            if item.filename.startswith(media_prefix):
                continue
            data = source.read(item.filename)
            if item.filename.startswith("word/_rels/") and item.filename.endswith(".rels"):
                text = data.decode("utf-8")
                data = relationship_media_pattern.sub("", text).encode("utf-8")
            target.writestr(item, data)
    temp_path.replace(path)


def _dedupe(values: list[str]) -> list[str]:
    seen: set[str] = set()
    result: list[str] = []
    for value in values:
        if value not in seen:
            seen.add(value)
            result.append(value)
    return result
